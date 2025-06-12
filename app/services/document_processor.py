import os
import uuid
import logging
import pymupdf
import docx
from PIL import Image
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional
import re

from app.models.schemas import FileType, DocumentChunk
from app.utils.config import settings

logger = logging.getLogger(__name__)

class SimpleTextSplitter:
    """Simple text splitter for French documents"""
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 150):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = ["\n\n", "\n", ". ", " "]
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        if not text:
            return []
        
        # Try to split by sentences first
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= self.chunk_size:
                current_chunk += sentence + " "
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + " "
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks

class FrenchDocumentProcessor:
    """Service for processing French cybersecurity documents with OCR support"""
    
    def __init__(self):
        self.ocr_reader = None  # Lazy initialization
        self.text_splitter = None
        self._initialize_services()
    
    def _initialize_services(self):
        """Initialize text splitting services (OCR initialized when needed)"""
        try:
            print("🔧 Initializing document processor...")
            
            # Initialize simple text splitter
            self.text_splitter = SimpleTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap
            )
            print("✅ Text splitter initialized for French documents")
            
        except Exception as e:
            print(f"❌ Error initializing document processor: {str(e)}")
            raise
    
    def _initialize_ocr(self):
        """Initialize OCR only when needed"""
        if self.ocr_reader is None:
            try:
                print("🔧 Initializing French OCR service (lazy loading)...")
                import easyocr
                self.ocr_reader = easyocr.Reader(['fr', 'en'], gpu=False)
                print("✅ French OCR service initialized")
            except Exception as e:
                print(f"❌ Error initializing OCR: {str(e)}")
                raise
    
    def detect_file_type(self, filename: str) -> FileType:
        """Detect file type from filename"""
        extension = Path(filename).suffix.lower()
        
        if extension == '.pdf':
            return FileType.PDF
        elif extension in ['.docx', '.doc']:
            return FileType.DOCX
        elif extension == '.txt':
            return FileType.TXT
        elif extension in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
            return FileType.IMAGE
        else:
            return FileType.TXT
    
    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            doc = pymupdf.open(file_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append(f"Page {page_num + 1}:\n{text}")
            
            doc.close()
            extracted_text = '\n\n'.join(text_content)
            print(f"📄 Extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            print(f"❌ Error extracting PDF text: {str(e)}")
            raise
    
    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_content.append('\n'.join(table_text))
            
            extracted_text = '\n\n'.join(text_content)
            print(f"📄 Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            print(f"❌ Error extracting DOCX text: {str(e)}")
            raise
    
    def extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    print(f"📄 Successfully read text file with {encoding} encoding")
                    return text
                except UnicodeDecodeError:
                    continue
            
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            print("⚠️ Used utf-8 with error replacement for text file")
            return text
            
        except Exception as e:
            print(f"❌ Error reading text file: {str(e)}")
            raise
    
    def extract_image_text(self, file_path: str) -> str:
        """Extract text from image using French OCR"""
        try:
            # Initialize OCR only when needed
            self._initialize_ocr()
            
            results = self.ocr_reader.readtext(file_path)
            
            extracted_text = []
            for (bbox, text, confidence) in results:
                if confidence > 0.5:
                    extracted_text.append(text)
            
            final_text = '\n'.join(extracted_text)
            print(f"📄 Extracted {len(final_text)} characters from image via OCR")
            return final_text
            
        except Exception as e:
            print(f"❌ Error extracting text from image: {str(e)}")
            raise
    
    def process_document(self, file_path: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process a document and extract text with metadata"""
        try:
            file_type = self.detect_file_type(filename)
            file_size = os.path.getsize(file_path)
            
            if file_type == FileType.PDF:
                extracted_text = self.extract_pdf_text(file_path)
            elif file_type == FileType.DOCX:
                extracted_text = self.extract_docx_text(file_path)
            elif file_type == FileType.TXT:
                extracted_text = self.extract_text_file(file_path)
            elif file_type == FileType.IMAGE:
                extracted_text = self.extract_image_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            metadata = {
                'filename': filename,
                'file_type': file_type,
                'file_size': file_size,
                'text_length': len(extracted_text),
                'word_count': len(extracted_text.split()) if extracted_text else 0,
                'language': 'french',
            }
            
            print(f"✅ Successfully processed {filename}: {metadata['text_length']} characters")
            return extracted_text, metadata
            
        except Exception as e:
            print(f"❌ Error processing document {filename}: {str(e)}")
            raise
    
    def create_chunks(self, text: str, document_id: str, filename: str) -> List[DocumentChunk]:
        """Split text into chunks optimized for French cybersecurity documents"""
        try:
            if not text or not text.strip():
                return []
            
            chunks = self.text_splitter.split_text(text)
            
            document_chunks = []
            for i, chunk_content in enumerate(chunks):
                if chunk_content.strip():
                    chunk = DocumentChunk(
                        chunk_id=f"{document_id}_chunk_{i}",
                        document_id=document_id,
                        content=chunk_content.strip(),
                        chunk_index=i,
                        metadata={
                            'chunk_length': len(chunk_content),
                            'word_count': len(chunk_content.split()),
                            'filename': filename  # <-- Add this line
                        }
                    )
                    document_chunks.append(chunk)
            
            print(f"📝 Created {len(document_chunks)} chunks for document {document_id}")
            return document_chunks
        
        except Exception as e:
            print(f"❌ Error creating chunks for document {document_id}: {str(e)}")
            return []
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        extension = Path(filename).suffix.lower().lstrip('.')
        return extension in settings.allowed_extensions
    
    def get_ocr_status(self) -> bool:
        """Check if OCR is initialized"""
        return self.ocr_reader is not None

# Global document processor instance
print("🚀 Creating French document processor...")
document_processor = FrenchDocumentProcessor()
print("✅ Document processor ready!")
